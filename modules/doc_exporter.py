# modules/doc_exporter.py
from docx import Document
from docx.shared import Inches
import os

def export_docx(image_path, output_docx):
    """
    Creates a DOCX file and inserts the graph image.
    
    Args:
        image_path (str): Path to the image file.
        output_docx (str): Path to save the DOCX file.
        
    Returns:
        str: The output path if successful.
    """
    try:
        doc = Document()
        doc.add_heading('VLM Graph Export', 0)
        
        doc.add_paragraph('Generated Graph:')
        
        # Add image, resizing to fit standard page width roughly
        doc.add_picture(image_path, width=Inches(6.0))
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(output_docx), exist_ok=True)
        
        doc.save(output_docx)
        return output_docx
    except Exception as e:
        raise RuntimeError(f"Failed to export DOCX: {str(e)}")
